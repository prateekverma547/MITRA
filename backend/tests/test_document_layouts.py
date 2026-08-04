"""What page layout does to extraction.

The eight files in `fixtures/documents/cv_layouts/` are the same CV for the same
fictional candidate, laid out eight ways, so any difference is caused by layout
alone. They are synthetic: they show what can happen, not how often it does.

**Only the DOCX cases are asserted here.** P3 and P4 extract in the wrong order
because `pypdf` reads the content stream as it was written, which is a different
cause needing a different fix. Pinning their current output would lock that bug
in, so this file deliberately does not.
"""

from pathlib import Path

import pytest

from blueprint.documents import DocumentError, extract_text

LAYOUTS = Path(__file__).parent / "fixtures" / "documents" / "cv_layouts"


def extracted(name: str) -> str:
    return extract_text(filename=name, data=(LAYOUTS / name).read_bytes())


# -- the identity block, wherever it lives -----------------------------------


def test_a_name_in_the_header_reaches_the_text():
    """The failure with a traced consequence.

    `document.paragraphs` does not include headers, so a CV that puts the name
    and contact details there lost all of it silently: the text began at
    EXPERIENCE. Downstream, generation asks the model for a name that is not in
    the text, gets none, and the interview opens with "Good afternoon." to
    somebody whose name we were given.
    """
    text = extracted("D4_header_contact.docx")

    assert "Arjun Mehta" in text
    assert "arjun.mehta@example.com" in text
    assert "Bengaluru, India" in text


def test_the_header_is_where_a_reader_expects_it():
    """At the top. It is the identity block, not a footnote."""
    text = extracted("D4_header_contact.docx")

    assert text.splitlines()[0].startswith("Arjun Mehta")
    assert text.index("Arjun Mehta") < text.index("EXPERIENCE")


def test_the_body_of_a_header_cv_is_still_all_there():
    """The header is added, nothing is displaced."""
    text = extracted("D4_header_contact.docx")

    # This sample has no education section: its identity block is in the header
    # and the body runs from EXPERIENCE to SKILLS.
    for claim in (
        "14 million transactions",
        "84 percent to 91 percent",
        "idempotency keys",
        "3,200 active merchants",
        "SQL, experiment design",
    ):
        assert claim in text


# -- nested tables -----------------------------------------------------------


def test_a_nested_table_cv_no_longer_fails():
    """`document.tables` returns only top-level tables. This CV had 99
    characters at the top level and 363 in two nested tables, so it fell under
    MIN_USABLE_CHARS and raised. That error was the only reason anyone noticed."""
    text = extracted("D3_nested_table.docx")

    assert len(text) > 400


def test_the_nested_content_is_actually_present():
    """Not just enough characters to clear the floor. The content itself."""
    text = extracted("D3_nested_table.docx")

    for claim in (
        "14 million transactions",
        "84 to 91 percent",
        "idempotency keys",
        "3,200 merchants",
        "Payments architecture",
        "B.E. Computer Science",
    ):
        assert claim in text


def test_a_nested_cv_reads_column_by_column():
    """Each cell is followed by what is inside it. Emitting the whole row first
    and the nested tables afterwards would separate every heading from its
    section."""
    text = extracted("D3_nested_table.docx")

    assert text.index("EXPERIENCE") < text.index("Own the card acquiring platform")
    assert text.index("Own the card acquiring platform") < text.index("SKILLS")
    assert text.index("SKILLS") < text.index("Payments architecture")


# -- layout tables versus data tables ----------------------------------------


def test_a_layout_table_keeps_all_of_its_content():
    text = extracted("D2_twocol_table.docx")

    for claim in (
        "ARJUN MEHTA",
        "14 million transactions",
        "idempotency keys",
        "3,200 merchants",
        "B.E. Computer Science",
        "arjun.mehta@example.com",
    ):
        assert claim in text


def test_a_layout_table_does_not_weld_a_heading_onto_a_sentence():
    """The pipe join is right for a data table and wrong for page layout, where
    it produced "...Introduced idempotency keys end to end. | EDUCATION": two
    separate things run together into one apparent sentence."""
    text = extracted("D2_twocol_table.docx")

    assert "|" not in text
    assert "end to end. | EDUCATION" not in text
    # The heading is its own line, with its own content under it.
    lines = text.splitlines()
    assert "EDUCATION" in lines
    assert lines[lines.index("EDUCATION") + 1] == "B.E. Computer Science"


def test_a_data_table_still_uses_pipes(tmp_path):
    """The case the fixtures do not cover, and the one a wrong rule would break.

    A table of values must be untouched by this change: each cell is one value,
    and the pipe is what keeps the pairing readable.
    """
    import docx

    document = docx.Document()
    table = document.add_table(rows=3, cols=2)
    for row, (skill, years) in enumerate(
        [("Skill", "Experience"), ("Python", "5 years"), ("SQL", "8 years")]
    ):
        table.rows[row].cells[0].text = skill
        table.rows[row].cells[1].text = years
    document.add_paragraph("Padding so the text clears the minimum length. " * 6)

    path = tmp_path / "data.docx"
    document.save(str(path))
    text = extract_text(filename="data.docx", data=path.read_bytes())

    assert "Python | 5 years" in text
    assert "SQL | 8 years" in text


# -- a plain document must not move ------------------------------------------


def test_a_plain_docx_is_untouched():
    """It works today. Everything here is additive to it."""
    text = extracted("D1_plain.docx")

    assert text.startswith("ARJUN MEHTA")
    assert "|" not in text
    for claim in (
        "Senior Product Manager, Payments. Bengaluru.",
        "14 million transactions",
        "arjun.mehta@example.com",
    ):
        assert claim in text


# -- the awkward shapes ------------------------------------------------------


def test_a_repeated_header_appears_once(tmp_path):
    """python-docx exposes one header per section rather than per page, but a
    document with several sections repeats it, either as its own copy or
    inherited. A three-section CV must not open with the name three times."""
    import docx
    from docx.enum.section import WD_SECTION

    document = docx.Document()
    document.sections[0].header.paragraphs[0].text = "Arjun Mehta | arjun@example.com"
    for _ in range(2):
        document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Body text long enough to clear the minimum. " * 8)

    path = tmp_path / "sections.docx"
    document.save(str(path))
    text = extract_text(filename="sections.docx", data=path.read_bytes())

    assert text.count("Arjun Mehta") == 1


def test_nesting_is_bounded(tmp_path):
    """Guarded rather than trusted to terminate. A generator gone wrong, or a
    hand-made file, must not walk forever."""
    import docx

    from blueprint.documents import MAX_TABLE_DEPTH

    document = docx.Document()
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    for depth in range(MAX_TABLE_DEPTH + 4):
        cell.text = f"level {depth}"
        inner = cell.add_table(rows=1, cols=1)
        cell = inner.rows[0].cells[0]
    cell.text = "the deepest level"
    document.add_paragraph("Padding so the text clears the minimum length. " * 6)

    path = tmp_path / "deep.docx"
    document.save(str(path))
    text = extract_text(filename="deep.docx", data=path.read_bytes())  # must not hang or raise

    assert "level 0" in text
    assert "the deepest level" not in text, "the depth guard did not stop the walk"


# -- the message a person reads ----------------------------------------------


def test_a_docx_is_not_told_about_scanned_pdfs(tmp_path):
    """It was. Wrong format, and useless as a next step."""
    import docx

    document = docx.Document()
    document.add_paragraph("too short")
    path = tmp_path / "thin.docx"
    document.save(str(path))

    with pytest.raises(DocumentError) as raised:
        extract_text(filename="thin.docx", data=path.read_bytes())

    assert "scanned PDF" not in str(raised.value)
    assert "text boxes" in str(raised.value)


def test_a_pdf_is_still_told_about_scanned_pdfs():
    with pytest.raises(DocumentError) as raised:
        extract_text(filename="scan.pdf", data=b"%PDF-1.4 not really a pdf")

    assert "Could not read PDF" in str(raised.value) or "scanned PDF" in str(raised.value)


# -- the PDF samples are not this change's business --------------------------


def test_the_clean_pdfs_still_extract_correctly():
    """P1 and P2 work today and this change must not have touched them."""
    for name in ("P1_single_column.pdf", "P2_twocol_ordered.pdf"):
        text = extracted(name)
        assert text.startswith("ARJUN MEHTA"), name
        assert "arjun.mehta@example.com" in text, name
        assert "idempotency keys" in text, name


def test_the_interleaved_pdfs_are_left_alone_deliberately():
    """P3 and P4 extract in the wrong order. That is `pypdf` reading the content
    stream as written, needs a library evaluation, and is not fixed here.

    Only that they still extract is asserted. Pinning the broken output would
    lock the bug in and make the real fix look like a regression.
    """
    for name in ("P3_twocol_interleaved.pdf", "P4_sidebar_builder.pdf"):
        text = extracted(name)
        assert text, name
