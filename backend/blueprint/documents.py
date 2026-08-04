"""Extract plain text from an uploaded JD or CV.

One adapter, three formats. Everything downstream — clarification chat, blueprint
generation — works on plain text and never learns what a PDF is.

Parsing quality matters more than it looks. A CV that extracts as garbage
produces a blueprint full of questions about things the candidate never claimed,
and the failure is silent: the blueprint is well-formed, just wrong. Hence the
sanity check on extracted text rather than trusting whatever comes back.
"""

import io
from pathlib import Path

MAX_UPLOAD_BYTES = 10 * 1024 * 1024

#: Below this, extraction almost certainly failed — a scanned PDF with no text
#: layer typically yields a handful of stray characters.
MIN_USABLE_CHARS = 200

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}


class DocumentError(ValueError):
    """Raised when a document cannot be read into usable text."""


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # noqa: BLE001
        raise DocumentError(f"Could not read PDF: {exc}") from exc
    return "\n\n".join(pages)


#: How far to follow tables nested inside table cells. Real CVs nest one or two
#: deep; a builder that goes further is malfunctioning and a hand-made file could
#: be pathological, so the walk is bounded rather than trusted to terminate.
MAX_TABLE_DEPTH = 5


def _row_text(cells: list[str]) -> str:
    """Join one row's cells, as a data table or as page layout.

    ` | ` is right for a data table, where each cell is one value: "Python | 5
    years" reads correctly and keeps the pairing. It is wrong for a table used to
    lay out a page, where it welds the sidebar onto the end of a body sentence:

        "Rebuilt retry logic after duplicate charges. Introduced idempotency
         keys end to end. | EDUCATION"

    Two separate things run together into one apparent sentence, and a heading
    left looking like the end of an achievement.

    **A cell containing a line break is holding a region of the page, not a
    value.** That is the difference, and it is a statement about what the two
    kinds of table are rather than a threshold fitted to these samples: one value
    is one line, a CV section is several. Decided per row, so a table whose cells
    are all single values keeps its pipes exactly as before and is untouched by
    this change.
    """
    return ("\n" if any("\n" in cell for cell in cells) else " | ").join(cells)


def _table_blocks(table, depth: int = 0) -> list[str]:
    """A row per block, following tables nested inside the cells.

    `document.tables` returns only top-level tables, so a CV whose sections are
    each their own table inside an outer layout table loses everything below the
    first level. One sample had 99 characters at the top level and 363 in two
    nested tables, and only the 99 were ever read.
    """
    if depth > MAX_TABLE_DEPTH:
        return []

    blocks: list[str] = []
    for row in table.rows:
        # A cell merged across columns is returned once per column it spans, so
        # a cell identical to the one before it is skipped. Compared by content
        # rather than by the underlying XML element, because lxml hands out
        # proxy objects and neither `is` nor `id()` reliably says whether two of
        # them are the same node.
        cells = []
        for cell in row.cells:
            if cells and cell.text == cells[-1].text and not cell.tables:
                continue
            cells.append(cell)

        texts = [c.text.strip() for c in cells if c.text.strip()]
        # A cell holding several lines, or holding a table of its own, is a
        # region of the page rather than a value. See `_row_text`.
        laid_out = any("\n" in t for t in texts) or any(c.tables for c in cells)

        if not laid_out:
            if texts:
                blocks.append(_row_text(texts))
            continue

        # Column by column, each cell followed by whatever it contains, which is
        # the order somebody reads the page in. Emitting the whole row first and
        # the nested tables afterwards would separate every section heading from
        # the section under it.
        for cell in cells:
            text = cell.text.strip()
            if text:
                blocks.append(text)
            for inner in cell.tables:
                blocks.extend(_table_blocks(inner, depth + 1))
    return blocks


def _part_text(part) -> list[str]:
    """Paragraphs and tables from a header or footer."""
    blocks = [p.text.strip() for p in part.paragraphs if p.text.strip()]
    for table in part.tables:
        blocks.extend(_table_blocks(table))
    return blocks


def _from_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise DocumentError(f"Could not read DOCX: {exc}") from exc

    # Headers and footers are not in `document.paragraphs`, so a CV that puts the
    # candidate's name and contact details in the header lost all of it and said
    # nothing: the text simply began at EXPERIENCE. Downstream, generation asks
    # the model for a name that is not there, gets none, and the interview opens
    # with "Good afternoon." to somebody whose name we were given.
    #
    # A header is the identity block and belongs at the top, where a reader
    # expects it. A footer is usually a page number, and is kept at the bottom
    # anyway: a stray "Page 1 of 2" costs a line of noise, while dropping footers
    # would lose the contact details some templates put there. The asymmetry
    # decides it.
    #
    # python-docx exposes one header per section rather than one per rendered
    # page, but a document with several sections repeats it, either as its own
    # copy or inherited through `is_linked_to_previous`. So identical blocks are
    # kept once: a three-section CV must not open with the name three times.
    headers: list[str] = []
    footers: list[str] = []
    for section in document.sections:
        for block in _part_text(section.header):
            if block not in headers:
                headers.append(block)
        for block in _part_text(section.footer):
            if block not in footers:
                footers.append(block)

    body = [p.text for p in document.paragraphs]
    # Plenty of real CVs lay everything out in tables; ignoring them loses the
    # entire work history.
    for table in document.tables:
        body.extend(_table_blocks(table))

    return "\n".join([*headers, *body, *footers])


def _from_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentError("Could not decode text file in utf-8, utf-16 or latin-1.")


def extract_text(*, filename: str, data: bytes) -> str:
    """Turn an uploaded document into plain text.

    Raises:
        DocumentError: unsupported type, unreadable file, or text so short that
            extraction evidently failed.
    """
    if not data:
        raise DocumentError("The uploaded file is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise DocumentError(
            f"File is {len(data) / 1_048_576:.1f} MB; the limit is "
            f"{MAX_UPLOAD_BYTES // 1_048_576} MB."
        )

    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise DocumentError(
            f"Unsupported file type '{suffix or filename}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_SUFFIXES))}."
        )

    if suffix == ".pdf":
        text = _from_pdf(data)
    elif suffix == ".docx":
        text = _from_docx(data)
    else:
        text = _from_text(data)

    text = normalise(text)

    if len(text) < MIN_USABLE_CHARS:
        raise DocumentError(
            f"Only {len(text)} characters of text could be extracted from "
            f"'{filename}'. {_why_so_little(suffix)}"
        )
    return text


def _why_so_little(suffix: str) -> str:
    """Advice that matches the file they actually uploaded.

    A DOCX uploader was being told their file might be a scanned PDF with no
    text layer, which is the wrong format and useless as a next step.
    """
    if suffix == ".pdf":
        return (
            "If this is a scanned PDF it has no text layer, so please upload a "
            "text-based version."
        )
    if suffix == ".docx":
        return (
            "Most of this document may be inside images or text boxes, which "
            "cannot be read. Please upload a version with the text as text, or "
            "save it as a PDF and upload that."
        )
    return "Please check the file and upload it again."


def normalise(text: str) -> str:
    """Collapse the whitespace damage that PDF extraction reliably produces."""
    lines = [" ".join(line.split()) for line in text.splitlines()]

    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if line:
            blank_run = 0
            cleaned.append(line)
        else:
            blank_run += 1
            if blank_run == 1:
                cleaned.append("")
    return "\n".join(cleaned).strip()
